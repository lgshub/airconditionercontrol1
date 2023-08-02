# encoding:utf-8
import gc
import pandas as pd
import numpy as np
import seaborn as sns
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import os
import logging
from PIL import ImageGrab
import xlwings as xw

logging.basicConfig(
    level=logging.INFO, format='%(levelname)s: %(message)s')


class HeatMap:
    def __init__(self, in_data_path, in_loc_file, out_heat_jg, out_heat_doc_file,
                 out_mid_data, roomiddic):  # 陕西  "./data/data-20210524-835/" './room_loc/'

        self.data_path = in_data_path  # 数据所在文件夹
        # self.data_read_time, self.JF_Name_ID = (in_data_path.split("/")[-1]).split("-")[1:]
        self.res_data_path = out_mid_data
        self.loc_file = in_loc_file
        self.processed_loc_file = out_mid_data + 'processed_loc.xlsx'
        self.plot_path = out_heat_jg
        self.out_heat_doc_file = out_heat_doc_file
        self.room_id_dic = roomiddic

    def process_loc(self):
        all_loc = pd.ExcelFile(self.loc_file)
        room_ids = all_loc.sheet_names  # 显示出读入excel文件中的表名字
        writer = pd.ExcelWriter(self.processed_loc_file)

        for room_id in room_ids:
            table = all_loc.parse(room_id)

            data = table.values
            data_ = pd.DataFrame(data[1:-1, 1:])
            if not self.room_id_dic:
                data_.to_excel(writer, sheet_name=room_id, index=None, header=None)
            else:
                data_.to_excel(writer, sheet_name=self.room_id_dic[room_id], index=None, header=None)

        writer.save()

    def get_room_id(self):  # 通过机柜物理定位表获得room_id
        all_loc = pd.ExcelFile(self.processed_loc_file)
        room_ids = all_loc.sheet_names
        return room_ids

    def get_temp_data(self):
        conv = {'room_id': str, 'rack_id': str, 'location': int,
                'temp': lambda x: -1 if x == "IDCNONE" else float(x)}
        data = pd.read_csv(self.data_path + "7.机柜温度数据.csv", converters=conv)
        return data

    def get_power_data(self):
        conv = {'room_id': str, 'dev_id': str, 'dev_type': int,
                'active_power': lambda x: -1 if x == "IDCNONE" else float(x),
                'ele_currnt': lambda x: -1 if x == "IDCNONE" else float(x),
                'voltage': lambda x: 1 if x == "IDCNONE" else float(x)}
        df_dldy = pd.read_csv(self.data_path + '5.电流电压数据.csv', converters=conv)
        return df_dldy

    def process_data(self, location=0):
        room_ID = self.get_room_id()
        data = self.get_temp_data()
        logging.info("***开始处理温度数据***")
        self.process_temp(data=data, room_ID=room_ID, location=location)
        logging.info("***开始释放温度数据***")
        del data
        gc.collect()

        df_dldy = self.get_power_data()
        logging.info("***开始处理功率数据***")
        self.process_power(data=df_dldy, room_ID=room_ID)
        logging.info("***开始释放功率数据***")
        del df_dldy
        gc.collect()

    def get_1h(self, data, usecols, keys=['time', 'room_id', 'acu_id']):
        data['time'] = data['time'].apply(lambda x: str(x)[:10])
        mydf = data.drop_duplicates(subset=keys,
                                    keep='first').reset_index(drop='True')
        mydf = mydf.sort_values(by='time')
        return mydf[usecols]

    def process_temp(self, data, room_ID, location):
        data = data.loc[data["room_id"].isin(room_ID)]
        usecols1 = ['time', 'room_id', 'rack_id', 'location', 'test_height', 'temp']
        data = self.get_1h(data, usecols1,
                           keys=['time', 'room_id', 'rack_id', 'location', 'test_height']
                           )
        data_location = data.loc[data["location"] == location]  # 冷通道 还是 热通道
        data_max = data_location.groupby(['time', 'room_id', 'rack_id', 'location'], as_index=False)['temp'].max()
        data_min = data_location.groupby(['time', 'room_id', 'rack_id', 'location'], as_index=False)['temp'].min()
        data_mean = data_location.groupby(['time', 'room_id', 'rack_id', 'location'], as_index=False)['temp'].mean()
        print("shape of data_max", data_max.shape)

        self.trans_JGdata(data=data_max, trans_col=['rack_id'], trans_val='temp',
                          file_str='max_temp')
        self.trans_JGdata(data=data_min, trans_col=['rack_id'], trans_val='temp',
                          file_str='min_temp')
        self.trans_JGdata(data=data_mean, trans_col=['rack_id'], trans_val='temp',
                          file_str='mean_temp')

    def process_power(self, data, room_ID):
        df_dldy = data.loc[data["room_id"].isin(room_ID)]
        usecols1 = ['time', 'room_id', 'dev_type', 'dev_id', 'ele_currnt', 'voltage', 'active_power']
        df_dldy = self.get_1h(df_dldy, usecols1,
                              keys=['time', 'room_id', 'dev_type', 'dev_id']
                              )
        # 部分电流出现负值波动，功率统一归为 0
        # df_dldy['active_power'] = np.where(df_dldy['active_power'] < 0, 0, df_dldy['active_power'])

        # 2021042817之前之间的功率单位是W,要转换为KW
        df_dldy['active_power'] = np.where(df_dldy['active_power'] > 10, df_dldy['active_power'] / 1000,
                                           df_dldy['active_power'])
        df_dldy['active_power'] = np.where(df_dldy['active_power'] == 0,
                                           df_dldy['ele_currnt'] * df_dldy['voltage'] / 1000,
                                           df_dldy['active_power'])

        # 分别将 空调1，机柜3拿出来
        df_dev_1 = df_dldy[df_dldy['dev_type'] == 1]
        df_dev_3 = df_dldy[df_dldy['dev_type'] == 3]
        # 将机柜A/B路相加
        # df_dev_3['dev_id'] = df_dev_3['dev_id'].apply(lambda x: str(x)[:9])
        # data2_1 = df_dev_3.groupby(['room_id','time', 'dev_id'],as_index=False)['active_power'].sum()

        # 将处理好的空调电量、机柜电量 长表转成符合现实位置的宽表
        # trans_JGdata(data=df_dev_1,trans_col=['dev_id'],trans_val='active_power',file_str='ACUpower')
        self.trans_JGdata(data=df_dev_3, trans_col=['dev_id'], trans_val='active_power',
                          file_str='ITpower')

        # 将空调功率存下
        for i in room_ID:
            data_x = df_dev_1.loc[(df_dev_1["room_id"] == i)]  # &(data1["time"] >='2020110500')

            datax1 = pd.pivot_table(data_x,
                                    index='time',
                                    columns=['dev_id'],
                                    values='active_power')

            # print(datax1.head(10))
            path = self.res_data_path
            if not os.path.exists(path):
                os.makedirs(path)
            datax1.to_csv(path + "/{0}{1}.csv".format('ACUpower', i))

    def trans_JGdata(self, data, trans_col=['rack_id'], trans_val='temp', file_str='max_temp'):
        room_list = sorted(data['room_id'].unique())
        all_loc = pd.ExcelFile(self.processed_loc_file)
        logging.info("***开始创建数据文件***")
        for i in room_list:

            loc = all_loc.parse(i, header=None)
            loc = loc.applymap(str)
            loc_matrix = loc.to_numpy().reshape(1, -1)[0].tolist()

            data_x = data.loc[(data["room_id"] == i)]  # &(data1["time"] >='2020110500')
            datax1 = pd.pivot_table(data_x, index='time', columns=trans_col, values=trans_val)

            rack_id_list = datax1.columns.tolist()

            res = pd.DataFrame(np.zeros((datax1.shape[0], len(loc_matrix))),
                               index=datax1.index, columns=loc_matrix)

            res = self.get_res(rack_id_list, loc_matrix, datax1, res)

            print(res.shape)
            path = self.res_data_path
            if not os.path.exists(path):
                os.makedirs(path)
            res.to_csv(path + "/{0}{1}.csv".format(file_str, i))
            print('Finish', i, '!')

    def get_res(self, rack_id_list, loc_matrix, data, res):
        for l in rack_id_list:
            if l in loc_matrix:
                res[l] = data[l]
        res["柱子"] = None
        res["nan"] = None  # 可能会导致res最后多加一列None，所以用下图清除
        while res.shape[1] != len(loc_matrix):
            res.drop(res.columns[len(res.columns) - 1], axis=1, inplace=True)
        return res

    def plot_headmap(self, time, room_ID="all", temp_min_val=18, temp_max_val=36):
        logging.info("***开始绘制热力图***")
        if room_ID == "all":
            room_ID = self.get_room_id()
        # process_data(path=path,room_ID=room_ID) #only need once

        # time = '2021050710'  # choose an hour from the period to plot heatmap

        # room_ID = ["2", "3", "4", "14", "15"]
        self.jg2plot(time2plot=time, room_ID=room_ID, fname='max_temp', chiname='机柜最高温度',
                     min_val=temp_min_val,
                     max_val=temp_max_val)
        # self.jg2plot(time2plot=time, room_ID=room_ID, fname='min_temp', chiname='机柜最低温度', min_val=18,
        #         max_val=36)
        # self.jg2plot(time2plot=time, room_ID=room_ID, fname='mean_temp', chiname='机柜平均温度',
        #         min_val=18, max_val=36)
        self.jg2plot(time2plot=time, room_ID=room_ID, fname='ITPower', chiname='机柜负载',
                     min_val=0,
                     max_val=4.0)
        logging.info("***热力图绘制完毕***")

    def get_label(self, data_map, name, rack_id, min_val, max_val):
        for i in range(data_map.shape[0]):
            for j in range(data_map.shape[1]):

                if data_map[i, j] is not None:
                    if data_map[i, j] == -1:
                        name[i, j] = "IDCNONE"
                    elif data_map[i, j] < 0:
                        name[i, j] = "负值异常"
                    elif data_map[i, j] == 0:
                        name[i, j] = "零值"
                    elif data_map[i, j] < min_val:
                        name[i, j] = "低温异常" + '(%.2f)' % data_map[i, j]
                    elif data_map[i, j] < max_val:
                        name[i, j] = ""
                    else:
                        name[i, j] = ""

        return name

    def IDC_heatmap(self, data_map, time, filename, title, room, loc, rack_id, min_val=10, max_val=32):
        path = self.plot_path + time
        if not os.path.exists(path):
            os.makedirs(path)
        sns.set()
        np.random.seed(0)
        name = loc.values
        label = self.get_label(data_map, name, rack_id, min_val, max_val)  # 数据名称，机柜名称，报警信息下限，报警信息上限
        plt.clf()
        plt.rcParams['font.sans-serif'] = ['SimHei']
        plt.rcParams['axes.unicode_minus'] = False
        # plt.figure(figsize=(9,6), dpi=100)
        ax = sns.heatmap(data_map,
                         center=(min_val + max_val) * 0.5,
                         vmin=min_val, vmax=max_val,
                         cmap="rainbow",
                         annot=label,
                         annot_kws={'size': 5, 'weight': 'normal'},
                         fmt="",
                         )

        if not os.path.exists(path):
            os.makedirs(path)
        plt.title('{0}月{1}日{2}点{3}--{4}'.format(time[4:6], time[6:8], time[8:], title, room))
        plt.savefig(path + "/{0}-{1}-{2}.png".format(filename, time, room), dpi=200)
        plt.clf()
        # plt.show()

    def jg2plot(self, time2plot='2020111708', room_ID=['HXXID3'],
                fname='max_temp', chiname='机柜温度', min_val=15, max_val=38, loc_path='./room_loc/陕西IDC平面图/'):
        # fname:res data的文件名前缀
        all_loc = pd.ExcelFile(self.processed_loc_file)
        for i in room_ID:
            # data = pd.read_csv('./res_data/{0}{1}.csv'.format(fname, i))
            # data = pd.read_csv('./res_data/软件园/' + time_data + '/{0}{1}.csv'.format(fname, i))
            data = pd.read_csv(self.res_data_path + '/{0}{1}.csv'.format(fname, i))

            # data.fillna(0.0, inplace=True)
            # data[data < 0] = 0

            # print(data["nan"])

            loc = all_loc.parse(i, header=None)
            loc = loc.applymap(str)
            loc_matrix = loc.to_numpy().reshape(1, -1)[0].tolist()

            # loc = pd.read_excel('./room_loc/{0}.xlsx'.format(i), header=None)
            data['time'] = data['time'].astype(str)

            mydata = data.loc[data['time'] == time2plot]
            mydata = mydata.to_numpy()[:, 1:].reshape(-1, loc.shape[0], loc.shape[1])
            mydata = mydata.astype(np.float32)
            self.IDC_heatmap(mydata[0], time2plot, fname, chiname, i, loc, loc_matrix, min_val, max_val)

    def generate_docs(self, time, room_ids, room_name=None):
        logging.info("***开始生成分析文档***")
        path = self.plot_path + time
        doc = Document()  # doc对象
        for room_id in room_ids:
            if room_name is None:
                text = doc.add_heading(room_id, 0)  # 添加文字
            else:
                text = doc.add_heading(room_name[room_id], 0)  # 添加文字
            max_image = path + "/max_temp-" + time + "-" + room_id + ".png"
            doc.add_picture(max_image, width=Inches(5.8))
            max_image = path + "/ITPower-" + time + "-" + room_id + ".png"
            doc.add_picture(max_image, width=Inches(5.8))
            doc.add_page_break()
        doc.save(self.out_heat_doc_file)
        logging.info("***分析文档生成完毕***")

    def generate_docs_with_topology(self, time, room_ids, room_name=None):
        logging.info("***开始生成分析文档***")
        path = self.plot_path + time
        self.excel_catch_screen(self.loc_file, room_ids, room_name=room_name)
        doc = Document()  # doc对象
        for room_id in room_ids:
            if room_name is None:
                text = doc.add_heading(room_id, 0)  # 添加文字
            else:
                text = doc.add_heading(room_name[room_id], 0)  # 添加文字
            # from docx.oxml.ns import qn
            # text.font.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
            loc_image = self.res_data_path + room_id + "_image.png"
            doc.add_picture(loc_image, width=Inches(6))
            max_image = path + "/max_temp-" + time + "-" + room_id + ".png"
            doc.add_picture(max_image, width=Inches(5.8))
            max_image = path + "/ITPower-" + time + "-" + room_id + ".png"
            doc.add_picture(max_image, width=Inches(5.8))
            doc.add_page_break()
        doc.save(self.out_heat_doc_file)
        logging.info("***分析文档生成完毕***")

    def excel_catch_screen(self, shot_excel, room_ids, room_name=None):
        app = xw.App(visible=True, add_book=False)  # 使用xlwings的app启动
        wb = app.books.open(shot_excel)  # 打开文件
        for shot_sheetname in room_ids:
            mid_img_path = self.res_data_path + shot_sheetname + "_image.png"
            if not room_name:
                sheet = wb.sheets(shot_sheetname)  # 选定sheet
            else:
                sheet = wb.sheets(room_name[shot_sheetname])
            all = sheet.used_range  # 获取有内容的range
            all.api.CopyPicture()  # 复制图片区域
            sheet.api.Paste()  # 粘贴
            img_name = 'data'
            pic = sheet.pictures[0]  # 当前图片
            pic.api.Copy()  # 复制图片
            img = ImageGrab.grabclipboard()  # 获取剪贴板的图片数据
            img.save(mid_img_path)  # 保存图片
            pic.delete()  # 删除sheet上的图片
        wb.close()  # 不保存，直接关闭
        app.quit()